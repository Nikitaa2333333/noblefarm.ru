import docx
from docx.opc.constants import RELATIONSHIP_TYPE as RT

def get_hyperlinks(paragraph):
    """
    Extract hyperlinks and their text from a paragraph
    """
    hyperlinks = []
    # Search for w:hyperlink elements
    for child in paragraph._element:
        if child.tag.endswith('hyperlink'):
            rId = child.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
            if rId:
                rel = paragraph.part.rels[rId]
                url = rel._target
                # Extract text inside the hyperlink
                text = "".join([node.text for node in child.iter() if node.tag.endswith('t') and node.text])
                hyperlinks.append((text, url))
    return hyperlinks

doc = docx.Document("Список публикаций.docx")
with open("raw_publications.txt", "w", encoding="utf-8") as f:
    f.write("=== PARAGRAPHS ===\n")
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text or para._element.xpath('.//w:hyperlink'):
            links = get_hyperlinks(para)
            link_str = ", ".join([f"[{t}]({u})" for t, u in links])
            f.write(f"P{idx}: {text} | LINKS: {link_str}\n")
            
    f.write("\n=== TABLES ===\n")
    for idx, table in enumerate(doc.tables):
        f.write(f"\nTABLE {idx}:\n")
        for r_idx, row in enumerate(table.rows):
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace('\n', ' ')
                cell_links = []
                for p in cell.paragraphs:
                    cell_links.extend(get_hyperlinks(p))
                link_str = ", ".join([f"[{t}]({u})" for t, u in cell_links])
                row_cells.append(f"{cell_text} (LINKS: {link_str})")
            f.write(f"  R{r_idx}: {row_cells}\n")
