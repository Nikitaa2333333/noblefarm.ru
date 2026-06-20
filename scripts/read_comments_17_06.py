import docx
import os

doc_path = r"c:\Users\User\Downloads\oooo_template\17.06.2026\Комментарии 17.06  часть 2.docx"
if not os.path.exists(doc_path):
    print(f"File not found: {doc_path}")
    exit(1)

doc = docx.Document(doc_path)
out_path = r"c:\Users\User\Downloads\oooo_template\scripts\comments_17_06_text.txt"

def get_hyperlinks(paragraph):
    hyperlinks = []
    for child in paragraph._element:
        if child.tag.endswith('hyperlink'):
            rId = child.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
            if rId:
                try:
                    rel = paragraph.part.rels[rId]
                    url = rel._target
                    text = "".join([node.text for node in child.iter() if node.tag.endswith('t') and node.text])
                    hyperlinks.append((text, url))
                except Exception:
                    pass
    return hyperlinks

with open(out_path, "w", encoding="utf-8") as f:
    f.write("=== PARAGRAPHS ===\n")
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        links = get_hyperlinks(para)
        link_str = ", ".join([f"[{t}]({u})" for t, u in links])
        if text or link_str:
            f.write(f"P{idx}: {text} | LINKS: {link_str}\n")
            
    f.write("\n=== TABLES ===\n")
    for idx, table in enumerate(doc.tables):
        f.write(f"\nTABLE {idx}:\n")
        for r_idx, row in enumerate(table.rows):
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace('\n', ' ')
                cell_links = []
                for p in cell.paragraphs:
                    cell_links.extend(get_hyperlinks(p))
                link_str = ", ".join([f"[{t}]({u})" for t, u in cell_links])
                row_cells.append(f"{cell_text} (LINKS: {link_str})")
            f.write(f"  R{r_idx}: {row_cells}\n")

print(f"Extraction complete. Output written to {out_path}")
